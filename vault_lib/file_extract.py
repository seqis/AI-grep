#!/usr/bin/env python3
"""
File Extraction - Extract content from various file types.

Walk a directory tree and extract readable content from:
- Text files (.py, .js, .md, .txt, .sh, .json, etc.)
- Word documents (.docx)
- Excel spreadsheets (.xlsx)
- PDF files (.pdf)

Binary and unsupported files are skipped with warnings.
"""

import fnmatch
import logging
import mimetypes
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Generator, Optional

# Configure module logger
logger = logging.getLogger(__name__)

# File type constants
TEXT_EXTENSIONS = {
    # Programming languages
    '.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.c', '.cpp', '.h', '.hpp',
    '.cs', '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.scala', '.r',
    # Web
    '.html', '.htm', '.css', '.scss', '.sass', '.less', '.vue', '.svelte',
    # Data/config
    '.json', '.yaml', '.yml', '.xml', '.toml', '.ini', '.cfg', '.conf',
    '.env', '.properties',
    # Documentation/text
    '.md', '.markdown', '.txt', '.rst', '.adoc', '.org', '.tex', '.rtf',
    # Scripts
    '.sh', '.bash', '.zsh', '.fish', '.ps1', '.bat', '.cmd',
    # Other
    '.sql', '.graphql', '.proto', '.csv', '.log', '.gitignore', '.dockerignore',
}

# Document types requiring special handling
DOCX_EXTENSIONS = {'.docx'}
XLSX_EXTENSIONS = {'.xlsx', '.xls'}
PDF_EXTENSIONS = {'.pdf'}

# Binary files to skip silently
BINARY_EXTENSIONS = {
    '.exe', '.dll', '.so', '.dylib', '.bin', '.dat',
    '.zip', '.tar', '.gz', '.bz2', '.xz', '.7z', '.rar',
    '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.ico', '.svg', '.webp',
    '.mp3', '.mp4', '.avi', '.mov', '.wav', '.flac', '.ogg',
    '.ttf', '.otf', '.woff', '.woff2', '.eot',
    '.pyc', '.pyo', '.class', '.o', '.obj',
    '.db', '.sqlite', '.sqlite3',
}

# Default exclude patterns
DEFAULT_EXCLUDE_PATTERNS = [
    '.*',  # Hidden files/directories
    '__pycache__',
    'node_modules',
    '.git',
    '.venv',
    'venv',
    '*.pyc',
    '*.pyo',
    '.DS_Store',
    'Thumbs.db',
]

# Text encodings to try (in order)
TEXT_ENCODINGS = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'ascii']


def detect_file_type(filepath: Path) -> str:
    """
    Detect the type of a file based on extension and content.

    Args:
        filepath: Path to the file

    Returns:
        File type string: 'text', 'docx', 'xlsx', 'pdf', 'binary', or 'unknown'
    """
    suffix = filepath.suffix.lower()

    # Check by extension first
    if suffix in TEXT_EXTENSIONS:
        return 'text'
    elif suffix in DOCX_EXTENSIONS:
        return 'docx'
    elif suffix in XLSX_EXTENSIONS:
        return 'xlsx'
    elif suffix in PDF_EXTENSIONS:
        return 'pdf'
    elif suffix in BINARY_EXTENSIONS:
        return 'binary'

    # Try to guess from mimetype
    mime_type, _ = mimetypes.guess_type(str(filepath))
    if mime_type:
        if mime_type.startswith('text/'):
            return 'text'
        elif mime_type.startswith('image/') or mime_type.startswith('audio/') or mime_type.startswith('video/'):
            return 'binary'
        elif 'wordprocessingml' in mime_type or mime_type == 'application/msword':
            return 'docx'
        elif 'spreadsheetml' in mime_type or mime_type == 'application/vnd.ms-excel':
            return 'xlsx'
        elif mime_type == 'application/pdf':
            return 'pdf'

    # Unknown - will be handled as text attempt
    return 'unknown'


def is_text_file(filepath: Path) -> bool:
    """
    Check if a file appears to be a text file.

    Uses extension and attempts to read the first few bytes to verify.

    Args:
        filepath: Path to the file

    Returns:
        True if the file appears to be text, False otherwise
    """
    file_type = detect_file_type(filepath)

    if file_type == 'text':
        return True
    elif file_type in ('binary', 'docx', 'xlsx', 'pdf'):
        return False

    # For unknown types, try to read and detect binary content
    try:
        with open(filepath, 'rb') as f:
            chunk = f.read(8192)

        # Check for null bytes (strong indicator of binary)
        if b'\x00' in chunk:
            return False

        # Try to decode as UTF-8
        try:
            chunk.decode('utf-8')
            return True
        except UnicodeDecodeError:
            # Try latin-1 (always succeeds but may be wrong)
            try:
                chunk.decode('latin-1')
                # Check for high ratio of printable characters
                printable = sum(1 for b in chunk if 32 <= b < 127 or b in (9, 10, 13))
                return printable / len(chunk) > 0.8 if chunk else True
            except Exception:
                return False

    except (OSError, IOError):
        return False


def _extract_text_file(filepath: Path) -> Optional[str]:
    """
    Extract content from a text file, trying multiple encodings.

    Args:
        filepath: Path to the text file

    Returns:
        File content as string, or None if extraction failed
    """
    for encoding in TEXT_ENCODINGS:
        try:
            content = filepath.read_text(encoding=encoding)
            return content
        except UnicodeDecodeError:
            continue
        except (OSError, IOError) as e:
            logger.warning(f"Failed to read {filepath}: {e}")
            return None

    logger.warning(f"Could not decode {filepath} with any known encoding")
    return None


def _extract_docx(filepath: Path) -> Optional[str]:
    """
    Extract content from a Word document (.docx).

    Extracts paragraphs and tables.

    Args:
        filepath: Path to the .docx file

    Returns:
        Extracted text content, or None if extraction failed
    """
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed. Install with: pip install python-docx")
        return None

    try:
        doc = Document(filepath)
        parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Extract tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(' | '.join(cells))
            if table_rows:
                parts.append('\n'.join(table_rows))

        return '\n\n'.join(parts)

    except Exception as e:
        logger.warning(f"Failed to extract from docx {filepath}: {e}")
        return None


def _extract_xlsx(filepath: Path) -> Optional[str]:
    """
    Extract content from an Excel spreadsheet (.xlsx).

    Extracts cell values from all sheets.

    Args:
        filepath: Path to the .xlsx file

    Returns:
        Extracted text content, or None if extraction failed
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        logger.warning("openpyxl not installed. Install with: pip install openpyxl")
        return None

    try:
        wb = load_workbook(filepath, read_only=True, data_only=True)
        parts = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_parts = [f"## Sheet: {sheet_name}"]

            rows = []
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                    else:
                        row_values.append('')
                if any(row_values):  # Only include non-empty rows
                    rows.append(' | '.join(row_values))

            if rows:
                sheet_parts.extend(rows)
                parts.append('\n'.join(sheet_parts))

        wb.close()
        return '\n\n'.join(parts)

    except Exception as e:
        logger.warning(f"Failed to extract from xlsx {filepath}: {e}")
        return None


def _extract_pdf(filepath: Path) -> Optional[str]:
    """
    Extract content from a PDF file.

    Uses PyPDF2 for extraction.

    Args:
        filepath: Path to the .pdf file

    Returns:
        Extracted text content, or None if extraction failed
    """
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        logger.warning("PyPDF2 not installed. Install with: pip install PyPDF2")
        return None

    try:
        reader = PdfReader(filepath)
        parts = []

        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                parts.append(f"## Page {i}\n{text.strip()}")

        return '\n\n'.join(parts)

    except Exception as e:
        logger.warning(f"Failed to extract from pdf {filepath}: {e}")
        return None


def extract_file(filepath: Path) -> Optional[dict]:
    """
    Extract content from a single file.

    Automatically detects file type and uses appropriate extraction method.

    Args:
        filepath: Path to the file to extract

    Returns:
        Dictionary with:
            - filepath: str (absolute path)
            - content: str (extracted content)
            - file_type: str ('text', 'docx', 'xlsx', 'pdf')
            - size: int (file size in bytes)
            - extracted_at: str (ISO timestamp)
            - encoding: str (for text files, the successful encoding)
        Returns None if extraction failed or file is binary/unsupported.
    """
    if not filepath.exists():
        logger.warning(f"File does not exist: {filepath}")
        return None

    if not filepath.is_file():
        logger.warning(f"Not a file: {filepath}")
        return None

    file_type = detect_file_type(filepath)
    content = None
    encoding = None

    if file_type == 'binary':
        logger.debug(f"Skipping binary file: {filepath}")
        return None

    elif file_type == 'text':
        # Try multiple encodings
        for enc in TEXT_ENCODINGS:
            try:
                content = filepath.read_text(encoding=enc)
                encoding = enc
                break
            except UnicodeDecodeError:
                continue
            except (OSError, IOError) as e:
                logger.warning(f"Failed to read {filepath}: {e}")
                return None

        if content is None:
            logger.warning(f"Could not decode {filepath} with any known encoding")
            return None

    elif file_type == 'docx':
        content = _extract_docx(filepath)

    elif file_type == 'xlsx':
        content = _extract_xlsx(filepath)

    elif file_type == 'pdf':
        content = _extract_pdf(filepath)

    elif file_type == 'unknown':
        # Try as text first
        if is_text_file(filepath):
            for enc in TEXT_ENCODINGS:
                try:
                    content = filepath.read_text(encoding=enc)
                    encoding = enc
                    file_type = 'text'
                    break
                except UnicodeDecodeError:
                    continue

        if content is None:
            logger.debug(f"Skipping unsupported file type: {filepath}")
            return None

    if content is None:
        return None

    try:
        file_size = filepath.stat().st_size
    except OSError:
        file_size = 0

    result = {
        'filepath': str(filepath.resolve()),
        'content': content,
        'file_type': file_type,
        'size': file_size,
        'extracted_at': datetime.now(timezone.utc).isoformat(),
    }

    if encoding:
        result['encoding'] = encoding

    return result


def _should_exclude(path: Path, exclude_patterns: list[str]) -> bool:
    """
    Check if a path should be excluded based on patterns.

    Args:
        path: Path to check
        exclude_patterns: List of glob patterns to exclude

    Returns:
        True if path should be excluded
    """
    name = path.name

    for pattern in exclude_patterns:
        # Match against filename
        if fnmatch.fnmatch(name, pattern):
            return True
        # Also check if any parent matches (for directory patterns)
        for parent in path.parents:
            if fnmatch.fnmatch(parent.name, pattern):
                return True

    return False


def walk_directory(
    root_path: Path,
    exclude_patterns: Optional[list[str]] = None,
    include_extensions: Optional[set[str]] = None,
) -> Generator[tuple[Path, str, dict], None, None]:
    """
    Walk a directory tree and extract content from all supported files.

    Args:
        root_path: Root directory to start walking
        exclude_patterns: List of glob patterns to exclude (default: hidden files,
                         __pycache__, node_modules, .git, etc.)
        include_extensions: If provided, only process files with these extensions.
                           Use lowercase with dot (e.g., {'.py', '.md'})

    Yields:
        Tuples of (filepath, content, metadata) where:
            - filepath: Path object
            - content: Extracted text content
            - metadata: dict with file_type, size, extracted_at, encoding (if text)

    Example:
        for filepath, content, meta in walk_directory(Path("/project")):
            print(f"{filepath}: {meta['file_type']}, {meta['size']} bytes")
    """
    if exclude_patterns is None:
        exclude_patterns = DEFAULT_EXCLUDE_PATTERNS

    root_path = Path(root_path).resolve()

    if not root_path.exists():
        logger.error(f"Directory does not exist: {root_path}")
        return

    if not root_path.is_dir():
        logger.error(f"Not a directory: {root_path}")
        return

    for item in root_path.rglob('*'):
        # Skip directories
        if item.is_dir():
            continue

        # Check exclusions
        if _should_exclude(item, exclude_patterns):
            logger.debug(f"Excluding: {item}")
            continue

        # Check extension filter
        if include_extensions is not None:
            if item.suffix.lower() not in include_extensions:
                continue

        # Try to extract
        try:
            result = extract_file(item)
            if result is not None:
                filepath = Path(result['filepath'])
                content = result['content']
                metadata = {
                    'file_type': result['file_type'],
                    'size': result['size'],
                    'extracted_at': result['extracted_at'],
                }
                if 'encoding' in result:
                    metadata['encoding'] = result['encoding']

                yield (filepath, content, metadata)

        except PermissionError:
            logger.warning(f"Permission denied: {item}")
        except Exception as e:
            logger.warning(f"Error processing {item}: {e}")


def get_supported_extensions() -> dict[str, set[str]]:
    """
    Get all supported file extensions by type.

    Returns:
        Dictionary mapping file type to set of extensions
    """
    return {
        'text': TEXT_EXTENSIONS.copy(),
        'docx': DOCX_EXTENSIONS.copy(),
        'xlsx': XLSX_EXTENSIONS.copy(),
        'pdf': PDF_EXTENSIONS.copy(),
        'binary': BINARY_EXTENSIONS.copy(),
    }


if __name__ == '__main__':
    # Simple test/demo
    import argparse

    parser = argparse.ArgumentParser(description='Extract content from files')
    parser.add_argument('path', nargs='?', default='.', help='File or directory to extract')
    parser.add_argument('-v', '--verbose', action='store_true', help='Verbose output')
    args = parser.parse_args()

    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format='%(levelname)s: %(message)s'
    )

    path = Path(args.path)

    if path.is_file():
        result = extract_file(path)
        if result:
            print(f"File: {result['filepath']}")
            print(f"Type: {result['file_type']}")
            print(f"Size: {result['size']} bytes")
            if 'encoding' in result:
                print(f"Encoding: {result['encoding']}")
            print(f"Content preview: {result['content'][:200]}...")
        else:
            print(f"Could not extract content from {path}")
            sys.exit(1)
    else:
        count = 0
        total_size = 0
        for filepath, content, meta in walk_directory(path):
            count += 1
            total_size += meta['size']
            print(f"{filepath.relative_to(path.resolve())}: {meta['file_type']}, {meta['size']} bytes")

        print(f"\nTotal: {count} files, {total_size:,} bytes")
